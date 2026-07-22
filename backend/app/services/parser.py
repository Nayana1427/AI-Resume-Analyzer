from pathlib import Path
import re

import fitz
from docx import Document


print("parser.py loaded")


class ResumeParser:

    # =========================================================
    # MAIN ENTRY POINT
    # =========================================================

    @staticmethod
    def extract_text(file_path: str) -> str:

        extension = Path(file_path).suffix.lower()

        if extension == ".pdf":
            return ResumeParser.extract_pdf(file_path)

        if extension == ".docx":
            return ResumeParser.extract_docx(file_path)

        raise ValueError(
            "Unsupported file format. Only PDF and DOCX are supported."
        )

    # =========================================================
    # CLEAN URL
    # =========================================================

    @staticmethod
    def _clean_url(url: str) -> str:

        if not url:
            return ""

        url = str(url).strip()

        # Remove common PDF/DOCX hyperlink prefixes
        if url.lower().startswith("mailto:"):
            return ""

        # Remove accidental surrounding characters
        url = url.strip(
            " \t\r\n<>[](){}\"'"
        )

        url = url.rstrip(
            ".,;:"
        )

        return url

    # =========================================================
    # CHECK WHETHER URL IS USEFUL
    # =========================================================

    @staticmethod
    def _is_relevant_url(url: str) -> bool:

        if not url:
            return False

        lower = url.lower()

        return (
            "linkedin.com/" in lower
            or "github.com/" in lower
        )

    # =========================================================
    # NORMALIZE SOCIAL URL
    # =========================================================

    @staticmethod
    def _normalize_social_url(url: str) -> str:

        url = ResumeParser._clean_url(url)

        if not url:
            return ""

        if not url.lower().startswith(
            ("http://", "https://")
        ):
            url = "https://" + url

        return url

    # =========================================================
    # REMOVE DUPLICATE URLS
    # =========================================================

    @staticmethod
    def _unique_urls(urls):

        result = []
        seen = set()

        for url in urls:

            url = ResumeParser._normalize_social_url(
                url
            )

            if not url:
                continue

            key = url.lower().rstrip("/")

            if key in seen:
                continue

            seen.add(key)
            result.append(url)

        return result

    # =========================================================
    # PDF
    # =========================================================

    @staticmethod
    def extract_pdf(file_path: str) -> str:

        document = fitz.open(file_path)

        text_parts = []
        hyperlinks = []

        try:

            for page in document:

                # ---------------------------------------------
                # Extract visible text
                # ---------------------------------------------

                page_text = page.get_text("text")

                if page_text:
                    text_parts.append(
                        page_text.strip()
                    )

                # ---------------------------------------------
                # Extract clickable hyperlink annotations
                # ---------------------------------------------

                try:

                    links = page.get_links()

                    for link in links:

                        uri = link.get("uri")

                        if not uri:
                            continue

                        uri = ResumeParser._clean_url(
                            uri
                        )

                        if ResumeParser._is_relevant_url(
                            uri
                        ):
                            hyperlinks.append(
                                uri
                            )

                except Exception as error:

                    print(
                        "PDF hyperlink extraction warning:",
                        error
                    )

        finally:

            document.close()

        # -----------------------------------------------------
        # Remove duplicate links
        # -----------------------------------------------------

        hyperlinks = ResumeParser._unique_urls(
            hyperlinks
        )

        # -----------------------------------------------------
        # Add URLs to extracted text
        #
        # InformationExtractor can now detect them.
        # -----------------------------------------------------

        if hyperlinks:

            text_parts.append(
                "\n".join(hyperlinks)
            )

        return "\n".join(
            part
            for part in text_parts
            if part
        ).strip()

    # =========================================================
    # DOCX HYPERLINK EXTRACTION
    # =========================================================

    @staticmethod
    def _extract_docx_hyperlinks(document):

        hyperlinks = []

        # -----------------------------------------------------
        # Paragraph hyperlinks
        # -----------------------------------------------------

        for paragraph in document.paragraphs:

            for hyperlink in paragraph._p.xpath(
                ".//w:hyperlink"
            ):

                relationship_id = hyperlink.get(
                    "{http://schemas.openxmlformats.org/"
                    "officeDocument/2006/relationships}id"
                )

                if not relationship_id:
                    continue

                relationship = (
                    document.part.rels.get(
                        relationship_id
                    )
                )

                if not relationship:
                    continue

                url = getattr(
                    relationship,
                    "target_ref",
                    ""
                )

                if (
                    url
                    and ResumeParser._is_relevant_url(
                        url
                    )
                ):
                    hyperlinks.append(
                        url
                    )

        # -----------------------------------------------------
        # Also inspect relationships directly.
        #
        # This catches hyperlinks that may not appear as normal
        # paragraph hyperlink elements.
        # -----------------------------------------------------

        try:

            for relationship in document.part.rels.values():

                url = getattr(
                    relationship,
                    "target_ref",
                    ""
                )

                if not url:
                    continue

                if ResumeParser._is_relevant_url(
                    url
                ):
                    hyperlinks.append(
                        url
                    )

        except Exception as error:

            print(
                "DOCX relationship extraction warning:",
                error
            )

        return ResumeParser._unique_urls(
            hyperlinks
        )

    # =========================================================
    # DOCX TABLE TEXT
    # =========================================================

    @staticmethod
    def _extract_docx_table_text(document):

        lines = []

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:
                        lines.append(value)

        return lines

    # =========================================================
    # DOCX
    # =========================================================

    @staticmethod
    def extract_docx(file_path: str) -> str:

        document = Document(file_path)

        text_parts = []

        # -----------------------------------------------------
        # Normal paragraph text
        # -----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        # -----------------------------------------------------
        # Text stored inside tables
        # -----------------------------------------------------

        table_text = (
            ResumeParser._extract_docx_table_text(
                document
            )
        )

        text_parts.extend(
            table_text
        )

        # -----------------------------------------------------
        # Hyperlinks
        # -----------------------------------------------------

        hyperlinks = (
            ResumeParser._extract_docx_hyperlinks(
                document
            )
        )

        if hyperlinks:

            text_parts.append(
                "\n".join(hyperlinks)
            )

        return "\n".join(
            part
            for part in text_parts
            if part
        ).strip()